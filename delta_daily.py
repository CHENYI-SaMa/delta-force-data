#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
三角洲行动每日数据抓取工具
数据来源: kkrb.net (三角洲行动一图流)

用法:
    python delta_daily.py           # 打印到屏幕
    python delta_daily.py -o        # 同时保存到txt文件
    python delta_daily.py --json    # 输出JSON格式
    python delta_daily.py --word    # 导出Word文档
    python delta_daily.py --word -o # 导出Word并保存
    python delta_daily.py --word --outdir "D:\\输出目录"  # 指定Word输出目录
"""

import urllib.request
import urllib.parse
import json
import sys
import os
from datetime import datetime, timedelta
from collections import defaultdict
import http.cookiejar

# ============================================================
# 配置区 - 在这里修改你要追踪的子弹
# ============================================================
TRACKED_AMMO = [
    ".357 Magnum FMJ",
    ".357 Magnum JHP",
    "玻纤柳叶箭矢",
    "9x39mm SP5",
    "12 Gauge独头 AP-20",
]

# 密码地点映射
PASSWORD_PLACES = {
    "db": "零号大坝",
    "cgxg": "长弓溪谷",
    "bks": "巴克什",
    "htjd": "航天基地",
    "cxjy": "潮汐监狱",
    "az3": "AZ3核电站",
    "az3r6": "AZ3彩六联动房",
}

# 制造台映射
WORKBENCH_NAMES = {
    "技术中心": "技术中心",
    "工作台": "工作台",
    "制药台": "制药台",
    "防具台": "防具台",
}

# ============================================================
# API 请求封装
# ============================================================
BASE_URL = "https://www.kkrb.net"


class KkrbAPI:
    """kkrb.net API 客户端"""

    def __init__(self):
        self.cj = http.cookiejar.CookieJar()
        self.opener = urllib.request.build_opener(
            urllib.request.HTTPCookieProcessor(self.cj)
        )
        self.opener.addheaders = [
            (
                "User-Agent",
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/131.0.0.0 Safari/537.36",
            ),
            ("Referer", f"{BASE_URL}/"),
            ("Origin", BASE_URL),
            ("X-Requested-With", "XMLHttpRequest"),
            ("Accept", "application/json, text/javascript, */*; q=0.01"),
            ("Accept-Language", "zh-CN,zh;q=0.9,en;q=0.8"),
        ]
        self._session_ready = False

    def _init_session(self):
        """初始化会话：访问主页 + 调用 getMenu 建立 session"""
        if self._session_ready:
            return
        self.opener.open(f"{BASE_URL}/")
        data = urllib.parse.urlencode({"globalData": "false"}).encode()
        req = urllib.request.Request(
            f"{BASE_URL}/getMenu", data=data, method="POST"
        )
        self.opener.open(req)
        self._session_ready = True

    def post(self, endpoint, params=None):
        """POST 请求 API"""
        self._init_session()
        if params is None:
            params = {}
        data = urllib.parse.urlencode(params).encode()
        req = urllib.request.Request(
            f"{BASE_URL}/{endpoint}", data=data, method="POST"
        )
        resp = self.opener.open(req, timeout=30)
        return json.loads(resp.read().decode("utf-8"))

    # ---- 各数据接口 ----

    def get_passwords(self):
        """获取今日密码"""
        result = self.post("getBonusDoorData")
        if result.get("code") != 1:
            raise Exception(f"密码接口返回错误: {result.get('msg', 'unknown')}")
        return result["data"]

    def get_manufacturing(self):
        """获取制造推荐"""
        result = self.post("getSwatProductData")
        if result.get("code") != 1:
            raise Exception(f"制造接口返回错误: {result.get('msg', 'unknown')}")
        return result.get("data", {}).get("cn", [])

    def get_exchange(self):
        """获取兑换推荐"""
        result = self.post("getTradeAmmoData")
        if result.get("code") != 1:
            raise Exception(f"兑换接口返回错误: {result.get('msg', 'unknown')}")
        return result.get("data", {}).get("cn", [])

    def get_ammo_prices(self):
        """获取所有子弹价格数据（含昨日最高/最低）"""
        result = self.post("getAmmoPLData", {"globalData": "false"})
        if result.get("code") != 1:
            raise Exception(f"子弹价格接口返回错误: {result.get('msg', 'unknown')}")
        return result.get("data", [])


# ============================================================
# 数据格式化
# ============================================================
def format_passwords(data):
    """格式化密码数据"""
    lines = []
    lines.append("【今日密码】")
    for key, name in PASSWORD_PLACES.items():
        if key in data:
            pwd = data[key].get("password", "未知")
            lines.append(f"  {name}: {pwd}")
    return "\n".join(lines)


def format_manufacturing(items):
    """格式化制造推荐 - 每个台子取利润最高的"""
    lines = []
    lines.append("【制造推荐】")

    # 按制造台分组
    by_place = defaultdict(list)
    for item in items:
        place = item.get("placeName", "未知")
        by_place[place].append(item)

    for place in ["技术中心", "工作台", "制药台", "防具台"]:
        if place in by_place:
            # 按利润排序取第一个
            sorted_items = sorted(
                by_place[place], key=lambda x: x.get("profit", 0), reverse=True
            )
            top = sorted_items[0]
            name = top.get("itemName", "?")
            profit = top.get("profit", 0)
            lines.append(f"  {place}: {name} (利润: {profit:,.0f})")

    return "\n".join(lines)


def format_exchange(items):
    """格式化兑换推荐 - 取收益前3"""
    lines = []
    lines.append("【兑换推荐】")

    sorted_items = sorted(items, key=lambda x: x.get("profit", 0), reverse=True)
    for i, item in enumerate(sorted_items[:3], 1):
        name = item.get("itemName", "?")
        profit = item.get("profit", 0)
        lines.append(f"  {i}. {name} (收益: {profit:,.0f})")

    return "\n".join(lines)


def format_ammo_prices(all_ammo, tracked_names):
    """格式化指定子弹的昨日最高/最低价"""
    lines = []
    lines.append("【指定子弹 昨日价格】")
    lines.append(f"  {'子弹名称':<22} {'昨日最高':>8} {'昨日最低':>8} {'当前价':>8}")
    lines.append(f"  {'-'*22} {'-'*8} {'-'*8} {'-'*8}")

    # 建立名称索引（支持模糊匹配）
    ammo_map = {}
    for item in all_ammo:
        ammo_map[item.get("itemName", "")] = item

    for name in tracked_names:
        # 精确匹配
        item = ammo_map.get(name)
        # 模糊匹配
        if not item:
            for key, val in ammo_map.items():
                if name.lower().replace(" ", "") in key.lower().replace(" ", ""):
                    item = val
                    name = key  # 用实际名称
                    break

        if item:
            high = item.get("yesterdayHighestPrice", 0)
            low = item.get("yesterdayLowestPrice", 0)
            curr = item.get("currectPrice", 0)
            lines.append(
                f"  {name:<22} {high:>8,} {low:>8,} {curr:>8,}"
            )
        else:
            lines.append(f"  {name:<22} {'未找到':>8}")

    return "\n".join(lines)


def format_report(api):
    """获取所有数据并格式化为完整报告"""
    today = datetime.now().strftime("%Y-%m-%d")
    lines = []
    lines.append("=" * 50)
    lines.append(f"  三角洲行动每日数据 - {today}")
    lines.append("=" * 50)
    lines.append("")

    # 1. 密码
    try:
        pwd_data = api.get_passwords()
        lines.append(format_passwords(pwd_data))
    except Exception as e:
        lines.append(f"【今日密码】\n  获取失败: {e}")
    lines.append("")

    # 2. 制造推荐
    try:
        mfg_data = api.get_manufacturing()
        lines.append(format_manufacturing(mfg_data))
    except Exception as e:
        lines.append(f"【制造推荐】\n  获取失败: {e}")
    lines.append("")

    # 3. 兑换推荐
    try:
        exchange_data = api.get_exchange()
        lines.append(format_exchange(exchange_data))
    except Exception as e:
        lines.append(f"【兑换推荐】\n  获取失败: {e}")
    lines.append("")

    # 4. 指定子弹价格
    try:
        ammo_data = api.get_ammo_prices()
        lines.append(format_ammo_prices(ammo_data, TRACKED_AMMO))
    except Exception as e:
        lines.append(f"【指定子弹 昨日价格】\n  获取失败: {e}")
    lines.append("")

    lines.append("-" * 50)
    lines.append(f"数据来源: kkrb.net | 抓取时间: {datetime.now().strftime('%H:%M:%S')}")

    return "\n".join(lines)


def format_json(api):
    """获取所有数据并返回JSON"""
    result = {
        "date": datetime.now().strftime("%Y-%m-%d"),
        "fetch_time": datetime.now().strftime("%H:%M:%S"),
        "passwords": {},
        "manufacturing": [],
        "exchange": [],
        "tracked_ammo": [],
    }

    try:
        pwd_data = api.get_passwords()
        for key, name in PASSWORD_PLACES.items():
            if key in pwd_data:
                result["passwords"][name] = pwd_data[key].get("password", "")
    except Exception as e:
        result["passwords"] = {"error": str(e)}

    try:
        mfg_data = api.get_manufacturing()
        by_place = defaultdict(list)
        for item in mfg_data:
            by_place[item.get("placeName", "")].append(item)
        for place in ["技术中心", "工作台", "制药台", "防具台"]:
            if place in by_place:
                top = max(by_place[place], key=lambda x: x.get("profit", 0))
                result["manufacturing"].append({
                    "place": place,
                    "item": top.get("itemName", ""),
                    "profit": top.get("profit", 0),
                })
    except Exception as e:
        result["manufacturing"] = [{"error": str(e)}]

    try:
        exchange_data = api.get_exchange()
        sorted_items = sorted(exchange_data, key=lambda x: x.get("profit", 0), reverse=True)
        for item in sorted_items[:3]:
            result["exchange"].append({
                "item": item.get("itemName", ""),
                "profit": item.get("profit", 0),
            })
    except Exception as e:
        result["exchange"] = [{"error": str(e)}]

    try:
        ammo_data = api.get_ammo_prices()
        ammo_map = {item.get("itemName", ""): item for item in ammo_data}
        for name in TRACKED_AMMO:
            item = ammo_map.get(name)
            if not item:
                for key, val in ammo_map.items():
                    if name.lower().replace(" ", "") in key.lower().replace(" ", ""):
                        item = val
                        name = key
                        break
            if item:
                result["tracked_ammo"].append({
                    "name": name,
                    "yesterday_high": item.get("yesterdayHighestPrice", 0),
                    "yesterday_low": item.get("yesterdayLowestPrice", 0),
                    "current_price": item.get("currectPrice", 0),
                })
            else:
                result["tracked_ammo"].append({"name": name, "error": "not found"})
    except Exception as e:
        result["tracked_ammo"] = [{"error": str(e)}]

    return json.dumps(result, ensure_ascii=False, indent=2)


# ============================================================
# 数据抓取（统一获取，供各格式化函数使用）
# ============================================================
def fetch_all_data(api):
    """统一获取所有数据，返回 dict"""
    data = {"date": datetime.now().strftime("%Y-%m-%d"),
            "fetch_time": datetime.now().strftime("%H:%M:%S")}

    # 密码
    try:
        pwd_raw = api.get_passwords()
        data["passwords"] = {}
        for key, name in PASSWORD_PLACES.items():
            if key in pwd_raw:
                data["passwords"][name] = pwd_raw[key].get("password", "")
    except Exception as e:
        data["passwords"] = None
        data["passwords_error"] = str(e)

    # 制造推荐
    try:
        mfg_raw = api.get_manufacturing()
        by_place = defaultdict(list)
        for item in mfg_raw:
            by_place[item.get("placeName", "")].append(item)
        data["manufacturing"] = []
        for place in ["技术中心", "工作台", "制药台", "防具台"]:
            if place in by_place:
                top = max(by_place[place], key=lambda x: x.get("profit", 0))
                data["manufacturing"].append({
                    "place": place,
                    "item": top.get("itemName", ""),
                    "profit": top.get("profit", 0),
                })
    except Exception as e:
        data["manufacturing"] = None
        data["manufacturing_error"] = str(e)

    # 兑换推荐
    try:
        ex_raw = api.get_exchange()
        sorted_items = sorted(ex_raw, key=lambda x: x.get("profit", 0), reverse=True)
        data["exchange"] = []
        for item in sorted_items[:3]:
            data["exchange"].append({
                "item": item.get("itemName", ""),
                "profit": item.get("profit", 0),
                "perCount": item.get("perCount", 1),
            })
    except Exception as e:
        data["exchange"] = None
        data["exchange_error"] = str(e)

    # 指定子弹价格
    try:
        ammo_raw = api.get_ammo_prices()
        ammo_map = {item.get("itemName", ""): item for item in ammo_raw}
        data["tracked_ammo"] = []
        for name in TRACKED_AMMO:
            item = ammo_map.get(name)
            if not item:
                for key, val in ammo_map.items():
                    if name.lower().replace(" ", "") in key.lower().replace(" ", ""):
                        item = val
                        name = key
                        break
            if item:
                data["tracked_ammo"].append({
                    "name": name,
                    "high": item.get("yesterdayHighestPrice", 0),
                    "low": item.get("yesterdayLowestPrice", 0),
                    "current": item.get("currectPrice", 0),
                })
            else:
                data["tracked_ammo"].append({"name": name, "error": "not found"})
    except Exception as e:
        data["tracked_ammo"] = None
        data["tracked_ammo_error"] = str(e)

    return data


# ============================================================
# Word 文档导出
# ============================================================
def export_word(data, output_path=None):
    """将数据导出为 Word 文档，格式参照用户原始消息风格"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        print("错误: 需要安装 python-docx 库")
        print("运行: pip install python-docx")
        return None

    doc = Document()

    # ---- 全局字体设置 ----
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # ---- 密码短名映射（匹配用户原始格式）----
    pwd_short = {
        "零号大坝": "大坝",
        "长弓溪谷": "长弓",
        "巴克什": "巴克",
        "航天基地": "航天",
        "AZ3核电站": "AZ3",
        "潮汐监狱": "监狱",
        "AZ3彩六联动房": "AZ3彩六联动房",
    }

    # ---- 标题 ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"三角洲行动每日数据  {data['date']}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    doc.add_paragraph()  # 空行

    # ---- 辅助函数：添加章节标题 ----
    def add_section_header(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        p.space_after = Pt(4)
        return p

    def add_content(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.bold = bold
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        return p

    # ============================================================
    # ➤ 今日密码
    # ============================================================
    add_section_header("➤今日密码:")

    if data.get("passwords"):
        pwds = data["passwords"]
        # 按用户原始格式：两个一行，用 ﹉ 分隔
        pwd_order = ["零号大坝", "长弓溪谷", "巴克什", "航天基地", "AZ3核电站", "潮汐监狱"]
        pwd_pairs = []
        for i in range(0, len(pwd_order), 2):
            pair = []
            for j in range(2):
                if i + j < len(pwd_order) and pwd_order[i + j] in pwds:
                    name = pwd_short.get(pwd_order[i + j], pwd_order[i + j])
                    pair.append(f"{name}:{pwds[pwd_order[i + j]]}")
            if pair:
                pwd_pairs.append("﹉".join(pair))

        for line in pwd_pairs:
            add_content(line)

        # AZ3彩六联动房 单独一行
        if "AZ3彩六联动房" in pwds:
            add_content(f"AZ3彩六联动房:{pwds['AZ3彩六联动房']}")
    else:
        add_content(f"获取失败: {data.get('passwords_error', '未知错误')}")

    doc.add_paragraph()  # 空行

    # ============================================================
    # ➤ 制造推荐
    # ============================================================
    add_section_header("➤制造推荐:")

    if data.get("manufacturing"):
        mfg = data["manufacturing"]
        # 每个台子单独一行，名字和利润分别一行（匹配用户截图）
        for item in mfg:
            add_content(f"{item['item']}")
            add_content(f"（利润:{item['profit']:,.0f}）")
    else:
        add_content(f"获取失败: {data.get('manufacturing_error', '未知错误')}")

    doc.add_paragraph()  # 空行

    # ============================================================
    # ➤ 兑换推荐
    # ============================================================
    add_section_header("➤兑换推荐:")
    add_content("战术部门:", bold=True)

    if data.get("exchange"):
        for item in data["exchange"]:
            add_content(f"{item['item']}")
            add_content(f"（收益:{item['profit']:,.0f}）")
    else:
        add_content(f"获取失败: {data.get('exchange_error', '未知错误')}")

    doc.add_paragraph()  # 空行

    # ============================================================
    # ➤ 指定子弹昨日价格（表格）
    # ============================================================
    add_section_header("➤指定子弹昨日价格")

    if data.get("tracked_ammo"):
        # 创建表格
        table = doc.add_table(rows=1 + len(data["tracked_ammo"]), cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        headers = ["类型", "最高", "最低"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            # 表头背景色
            from docx.oxml import OxmlElement
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "D9E2F3")
            cell._element.get_or_add_tcPr().append(shading)

        # 数据行
        for row_idx, item in enumerate(data["tracked_ammo"], 1):
            name = item.get("name", "?")
            if "error" in item:
                table.rows[row_idx].cells[0].text = name
                table.rows[row_idx].cells[1].text = "未找到"
                table.rows[row_idx].cells[2].text = ""
            else:
                high = item.get("high", 0)
                low = item.get("low", 0)
                for col_idx, value in enumerate([name, f"{high:,}", f"{low:,}"]):
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    if col_idx > 0:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(value)
                    run.font.size = Pt(11)
                    run.font.name = "微软雅黑"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    else:
        add_content(f"获取失败: {data.get('tracked_ammo_error', '未知错误')}")

    doc.add_paragraph()  # 空行

    # ---- 带横线分隔的纯文本版本 ----
    add_section_header("➤指定子弹昨日价格（横线版）")
    if data.get("tracked_ammo"):
        # 表头
        add_content("﹉﹉类型﹉﹉﹉﹉﹉﹉﹉﹉最高﹉﹉﹉最低")
        for item in data["tracked_ammo"]:
            if "error" in item:
                add_content(f"{item['name']}---------- 未找到")
            else:
                name = item.get("name", "?")
                high = item.get("high", 0)
                low = item.get("low", 0)
                add_content(f"{name}----------{high:,}-------{low:,}")

    doc.add_paragraph()  # 空行

    # ---- 页脚信息 ----
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(f"数据来源: kkrb.net  |  抓取时间: {data['fetch_time']}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # ---- 保存 ----
    if output_path is None:
        script_dir = os.path.dirname(os.path.abspath(__file__))
        output_path = os.path.join(script_dir, f"三角洲行动每日数据_{data['date']}.docx")

    doc.save(output_path)
    return output_path


# ============================================================
# 主程序
# ============================================================
def format_report_from_data(data):
    """从已抓取的数据生成文本报告（避免重复API请求）"""
    lines = []
    lines.append("=" * 50)
    lines.append(f"  三角洲行动每日数据 - {data['date']}")
    lines.append("=" * 50)
    lines.append("")

    # 密码
    lines.append("【今日密码】")
    if data.get("passwords"):
        for name, pwd in data["passwords"].items():
            lines.append(f"  {name}: {pwd}")
    else:
        lines.append(f"  获取失败: {data.get('passwords_error', '未知')}")
    lines.append("")

    # 制造
    lines.append("【制造推荐】")
    if data.get("manufacturing"):
        for item in data["manufacturing"]:
            lines.append(f"  {item['place']}: {item['item']} (利润: {item['profit']:,.0f})")
    else:
        lines.append(f"  获取失败: {data.get('manufacturing_error', '未知')}")
    lines.append("")

    # 兑换
    lines.append("【兑换推荐】")
    if data.get("exchange"):
        for i, item in enumerate(data["exchange"], 1):
            lines.append(f"  {i}. {item['item']} (收益: {item['profit']:,.0f})")
    else:
        lines.append(f"  获取失败: {data.get('exchange_error', '未知')}")
    lines.append("")

    # 子弹价格
    lines.append("【指定子弹 昨日价格】")
    lines.append(f"  {'子弹名称':<22} {'昨日最高':>8} {'昨日最低':>8} {'当前价':>8}")
    lines.append(f"  {'-'*22} {'-'*8} {'-'*8} {'-'*8}")
    if data.get("tracked_ammo"):
        for item in data["tracked_ammo"]:
            if "error" in item:
                lines.append(f"  {item['name']:<22} {'未找到':>8}")
            else:
                lines.append(f"  {item['name']:<22} {item['high']:>8,} {item['low']:>8,} {item['current']:>8,}")
    else:
        lines.append(f"  获取失败: {data.get('tracked_ammo_error', '未知')}")
    lines.append("")

    lines.append("-" * 50)
    lines.append(f"数据来源: kkrb.net | 抓取时间: {data['fetch_time']}")

    return "\n".join(lines)


def push_to_github(data_path):
    """推送 data.json 到 GitHub 更新网页。需要代理和 token 文件，失败静默跳过。"""
    import base64
    import socket

    SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
    TOKEN_FILE = os.path.join(SCRIPT_DIR, "delta-force-cloud", "github_token.txt")
    REPO = "CHENYI-SaMa/delta-force-data"

    # 检查 token 文件
    if not os.path.exists(TOKEN_FILE):
        return False, "no token file"
    with open(TOKEN_FILE, "r") as f:
        token = f.read().strip()
    if not token:
        return False, "empty token"

    # 检测代理端口（Clash/v2ray 常用端口）
    proxy_port = None
    for port in [7893, 7890, 7891, 7892, 1080, 10809]:
        try:
            s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            s.settimeout(0.5)
            if s.connect_ex(("127.0.0.1", port)) == 0:
                proxy_port = port
            s.close()
            if proxy_port:
                break
        except Exception:
            pass

    if not proxy_port:
        return False, "no proxy (accelerator not running)"

    proxy = f"http://127.0.0.1:{proxy_port}"
    proxy_handler = urllib.request.ProxyHandler({"https": proxy, "http": proxy})
    gh_opener = urllib.request.build_opener(proxy_handler)

    try:
        # 获取远程文件 sha
        req = urllib.request.Request(
            f"https://api.github.com/repos/{REPO}/contents/data.json",
            headers={
                "Authorization": f"Bearer {token}",
                "Accept": "application/vnd.github+json",
            },
        )
        resp = gh_opener.open(req, timeout=15)
        remote_data = json.loads(resp.read().decode("utf-8"))
        sha = remote_data["sha"]

        # 读取本地 data.json 并 base64 编码
        with open(data_path, "rb") as f:
            content = base64.b64encode(f.read()).decode()

        # 推送更新
        body = json.dumps(
            {
                "message": f"Auto-update from local: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
                "content": content,
                "sha": sha,
            }
        ).encode()
        req2 = urllib.request.Request(
            f"https://api.github.com/repos/{REPO}/contents/data.json",
            data=body,
            headers={
                "Authorization": f"Bearer {token}",
                "Accept": "application/vnd.github+json",
                "Content-Type": "application/json",
            },
            method="PUT",
        )
        resp2 = gh_opener.open(req2, timeout=15)
        result = json.loads(resp2.read().decode("utf-8"))
        return True, result["commit"]["sha"][:8]
    except Exception as e:
        return False, str(e)


def format_wxpusher_message(data):
    """格式化数据为 WxPusher 微信推送消息"""
    lines = []
    lines.append(f"三角洲行动每日数据 {data['date']}")
    lines.append("")

    # 密码短名映射
    pwd_short = {
        "零号大坝": "大坝",
        "长弓溪谷": "长弓",
        "巴克什": "巴克",
        "航天基地": "航天",
        "AZ3核电站": "AZ3",
        "潮汐监狱": "监狱",
        "AZ3彩六联动房": "AZ3彩六",
    }

    # 密码
    lines.append("➤今日密码")
    if data.get("passwords"):
        pwd_order = ["零号大坝", "长弓溪谷", "巴克什", "航天基地", "AZ3核电站", "潮汐监狱", "AZ3彩六联动房"]
        for name in pwd_order:
            if name in data["passwords"]:
                short = pwd_short.get(name, name)
                lines.append(f"  {short} {data['passwords'][name]}")
    else:
        lines.append("  获取失败")
    lines.append("")

    # 制造推荐
    lines.append("➤制造推荐")
    if data.get("manufacturing"):
        for item in data["manufacturing"]:
            lines.append(f"  {item['item']}")
            lines.append(f"  (利润:{item['profit']:,.0f})")
    else:
        lines.append("  获取失败")
    lines.append("")

    # 兑换推荐
    lines.append("➤兑换推荐")
    if data.get("exchange"):
        for item in data["exchange"]:
            lines.append(f"  {item['item']}")
            lines.append(f"  (收益:{item['profit']:,.0f})")
    else:
        lines.append("  获取失败")
    lines.append("")

    # 子弹价格
    lines.append("➤子弹昨日价格")
    if data.get("tracked_ammo"):
        for item in data["tracked_ammo"]:
            if "error" in item:
                lines.append(f"  {item['name']} 未找到")
            else:
                lines.append(f"  {item['name']}")
                lines.append(f"    最高 {item['high']:,} / 最低 {item['low']:,}")
    else:
        lines.append("  获取失败")
    lines.append("")
    lines.append(f"kkrb.net | {data['fetch_time']}")

    return "\n".join(lines)


def push_wxpusher(data):
    """通过 WxPusher 推送数据到微信。失败静默跳过。"""
    SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
    CONFIG_FILE = os.path.join(SCRIPT_DIR, "delta-force-cloud", "wxpusher_config.json")

    if not os.path.exists(CONFIG_FILE):
        return False, "no config file"
    try:
        with open(CONFIG_FILE, "r", encoding="utf-8") as f:
            config = json.load(f)
    except Exception:
        return False, "config file invalid"

    app_token = config.get("app_token", "")
    uids = config.get("uids", [])

    if not app_token or not uids:
        return False, "missing app_token or uids"

    content = format_wxpusher_message(data)
    summary = f"三角洲每日数据 {data['date']}"

    body = json.dumps({
        "appToken": app_token,
        "content": content,
        "summary": summary,
        "contentType": 1,
        "uids": uids,
    }).encode("utf-8")

    try:
        req = urllib.request.Request(
            "https://wxpusher.zjiecode.com/api/send/message",
            data=body,
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        resp = urllib.request.urlopen(req, timeout=15)
        result = json.loads(resp.read().decode("utf-8"))
        if result.get("code") == 1000:
            return True, "ok"
        else:
            return False, result.get("msg", "unknown error")
    except Exception as e:
        return False, str(e)


def main():
    args = sys.argv[1:]
    save_to_file = "-o" in args or "--output" in args
    json_mode = "--json" in args
    word_mode = "--word" in args

    # 解析 --outdir 参数
    outdir = None
    if "--outdir" in args:
        idx = args.index("--outdir")
        if idx + 1 < len(args):
            outdir = args[idx + 1]

    api = KkrbAPI()

    if word_mode:
        # Word 导出模式：抓取数据 -> 导出Word -> 打印文本摘要
        data = fetch_all_data(api)
        # 构造输出路径
        output_path = None
        if outdir:
            os.makedirs(outdir, exist_ok=True)
            output_path = os.path.join(outdir, f"三角洲行动每日数据_{data['date']}.docx")
        path = export_word(data, output_path)
        # 同时打印文本到屏幕
        text_output = format_report_from_data(data)
        print(text_output)
        if path:
            print(f"\nWord 文档已导出: {path}")
        # 写入运行日志（方便确认定时任务是否正常）
        if outdir:
            log_path = os.path.join(outdir, "run_log.txt")
            with open(log_path, "a", encoding="utf-8") as f:
                f.write(f"\n{'='*50}\n")
                f.write(f"运行时间: {data['date']} {data['fetch_time']}\n")
                f.write(text_output + "\n")
                if path:
                    f.write(f"Word: {path}\n")
                f.write(f"状态: {'成功' if path else '失败'}\n")
            # 同时输出 data.json 供手机网页读取
            json_path = os.path.join(outdir, "data.json")
            with open(json_path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            # 推送到 GitHub 更新网页
            ok, msg = push_to_github(json_path)
            if ok:
                print(f"GitHub 网页已更新 (commit: {msg})")
            else:
                print(f"GitHub 推送跳过: {msg}")
            # 推送到微信 (WxPusher)
            ok, msg = push_wxpusher(data)
            if ok:
                print("微信推送成功")
            else:
                print(f"微信推送跳过: {msg}")
        return

    if json_mode:
        output = format_json(api)
    else:
        output = format_report(api)

    print(output)

    if save_to_file:
        script_dir = os.path.dirname(os.path.abspath(__file__))
        today = datetime.now().strftime("%Y-%m-%d")
        ext = "json" if json_mode else "txt"
        filename = os.path.join(script_dir, f"delta_data_{today}.{ext}")
        with open(filename, "w", encoding="utf-8") as f:
            f.write(output)
        print(f"\n已保存到: {filename}")


if __name__ == "__main__":
    main()
